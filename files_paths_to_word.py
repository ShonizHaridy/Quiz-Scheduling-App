import os
from docx import Document

def write_files_to_word(directory, doc):
    # Walk through the directory
    for root, dirs, files in os.walk(directory):
        for file in files:
            # Create the full file path
            file_path = os.path.join(root, file)

            # Add file path to the Word document
            doc.add_paragraph(f"File Path: {file_path}")

def main():
    # Specify the directory you want to scan
    directory = input("Enter the directory path: ")

    # Create a new Word document
    doc = Document()
    doc.add_heading('File Paths', level=1)

    # Write the file paths to the document
    write_files_to_word(directory, doc)

    # Save the document
    output_file = os.path.join(directory, 'file_paths.docx')
    doc.save(output_file)
    print(f"Document saved as: {output_file}")

if __name__ == "__main__":
    main()
